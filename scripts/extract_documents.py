# extract_documents.py (Layer 2 upgrade)
import os
import json
from pathlib import Path

# PDF
import PyPDF2
from pdf2image import convert_from_path
import pytesseract
import re

# Word
import docx

# Excel / CSV
import openpyxl
import csv

# -------------------------------
# CONFIGURATION
# -------------------------------
BASE_DIR = Path(__file__).resolve().parent.parent
DATA_DIR = BASE_DIR / "data"
OUTPUT_DIR = BASE_DIR / "structured_docs"
OUTPUT_DIR.mkdir(exist_ok=True)

CHUNK_SIZE = 1000  # characters

# -------------------------------
# HELPER FUNCTIONS
# -------------------------------
def detect_heading_pdf(text_line):
    line = text_line.strip()
    if line.isupper() or line.endswith(":"):
        return line
    return None

def detect_heading_docx(paragraph):
    style = paragraph.style.name
    if "Heading" in style:
        return paragraph.text.strip()
    return None

def chunk_text(text, source_name, page_num, section=None):
    chunks = []
    text = text.strip().replace("\n", " ")
    for i in range(0, len(text), CHUNK_SIZE):
        chunk = text[i:i+CHUNK_SIZE]
        chunks.append({
            "source": source_name,
            "page": page_num,
            "section": section,
            "text": chunk
        })
    return chunks

def extract_tables_from_docx(doc):
    tables_chunks = []
    for table_num, table in enumerate(doc.tables, start=1):
        for row_num, row in enumerate(table.rows, start=1):
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                tables_chunks.extend(chunk_text(row_text, f"Table-{table_num}", row_num))
    return tables_chunks

# -------------------------------
# EXTRACTION FUNCTIONS
# -------------------------------
def extract_pdf(file_path):
    text_chunks = []
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                section = None
                if text and text.strip():
                    # detect headings
                    for line in text.split("\n"):
                        heading = detect_heading_pdf(line)
                        if heading:
                            section = heading
                    # detect tables (simple heuristic: multiple | or tab spacing)
                    table_lines = [l for l in text.split("\n") if re.search(r"\s{2,}|\|", l)]
                    for t_line in table_lines:
                        text_chunks.extend(chunk_text(t_line, file_path.name, page_num, section))
                    text_chunks.extend(chunk_text(text, file_path.name, page_num, section))
                else:
                    # OCR fallback
                    images = convert_from_path(file_path)
                    for i, image in enumerate(images):
                        ocr_text = pytesseract.image_to_string(image)
                        text_chunks.extend(chunk_text(ocr_text, file_path.name, i+1))
    except Exception as e:
        print(f"[ERROR] PDF extraction failed for {file_path}: {e}")
    return text_chunks

def extract_docx(file_path):
    text_chunks = []
    try:
        doc = docx.Document(file_path)
        section = None
        for i, para in enumerate(doc.paragraphs, start=1):
            heading = detect_heading_docx(para)
            if heading:
                section = heading
            if para.text.strip():
                text_chunks.extend(chunk_text(para.text, file_path.name, i, section))
        # extract tables
        text_chunks.extend(extract_tables_from_docx(doc))
    except Exception as e:
        print(f"[ERROR] DOCX extraction failed for {file_path}: {e}")
    return text_chunks

def extract_xlsx(file_path):
    text_chunks = []
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            for row_num, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                row_text = " | ".join([str(cell) for cell in row if cell is not None])
                if row_text.strip():
                    text_chunks.extend(chunk_text(row_text, file_path.name, row_num, section=sheet_name))
    except Exception as e:
        print(f"[ERROR] XLSX extraction failed for {file_path}: {e}")
    return text_chunks

def extract_csv(file_path):
    text_chunks = []
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            for row_num, row in enumerate(reader, start=1):
                row_text = " | ".join([str(cell) for cell in row if cell])
                if row_text.strip():
                    text_chunks.extend(chunk_text(row_text, file_path.name, row_num))
    except Exception as e:
        print(f"[ERROR] CSV extraction failed for {file_path}: {e}")
    return text_chunks

def process_file(file_path):
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(file_path)
    elif suffix == ".docx":
        return extract_docx(file_path)
    elif suffix in [".xlsx", ".xls"]:
        return extract_xlsx(file_path)
    elif suffix == ".csv":
        return extract_csv(file_path)
    else:
        print(f"[WARNING] Unsupported file type: {file_path}")
        return []

# -------------------------------
# MAIN PIPELINE
# -------------------------------
def main():
    print("[INFO] Starting document processing...")
    for file_path in DATA_DIR.iterdir():
        if file_path.is_file():
            print(f"[INFO] Processing {file_path.name}...")
            chunks = process_file(file_path)
            if chunks:
                output_file = OUTPUT_DIR / f"structured_{file_path.stem}.json"
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(chunks, f, indent=2, ensure_ascii=False)
                print(f"[INFO] Saved {len(chunks)} chunks to {output_file.name}")
            else:
                print(f"[INFO] No chunks extracted from {file_path.name}")
    print("[INFO] Document processing complete.")

if __name__ == "__main__":
    main()
